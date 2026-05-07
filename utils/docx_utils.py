"""DOCX parsing utilities for VerdictAI Document Intelligence layer (L1).

Parses Microsoft Word (.docx) files using ``python-docx``. DOCX files
carry embedded text (no OCR needed), so they go straight through the
digital-text path in L1 — the result shape mirrors :func:`parse_pdf` so
the rest of the pipeline needs zero branching logic.

Design choices:
- ``page_count`` is always 1. DOCX has no reliable page concept (the
  rendered page count depends on the Word version, printer, and
  page-layout settings). We present the whole document as a single
  "logical page" so the downstream schema (pages, word_objects) still
  works unchanged.
- Tables are flattened into tab-separated text and appended after the
  paragraph text. This preserves the content for the evidence
  extractor's regex pass without needing a separate table-aware path
  for DOCX.
- ``is_scanned`` is always False. DOCX always has extractable text.
"""

from __future__ import annotations

import logging
import os

logger = logging.getLogger(__name__)


def parse_docx(file_path: str) -> dict:
    """Parse a DOCX file and return a pdf_utils-compatible shape.

    Args:
        file_path: Path to the .docx file.

    Returns:
        On success::

            {
              "page_count": 1,
              "pages": [{"page_number": 1, "text": str, "has_tables": bool}],
              "is_scanned": False,
            }

        On failure::

            {"error": True, "message": str, "page_count": 0,
             "pages": [], "is_scanned": False}
    """
    if not os.path.exists(file_path):
        return {
            "error": True,
            "message": f"File not found: {file_path}",
            "page_count": 0,
            "pages": [],
            "is_scanned": False,
        }

    try:
        # Import lazily so tests that don't hit DOCX don't pay the
        # import cost and missing-dep cases produce a clean error.
        from docx import Document  # type: ignore
    except ImportError as exc:
        return {
            "error": True,
            "message": f"python-docx not installed: {exc}",
            "page_count": 0,
            "pages": [],
            "is_scanned": False,
        }

    try:
        doc = Document(file_path)
    except Exception as exc:
        return {
            "error": True,
            "message": (
                f"Failed to open DOCX '{file_path}': "
                f"{type(exc).__name__}: {exc}"
            ),
            "page_count": 0,
            "pages": [],
            "is_scanned": False,
        }

    text_parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text and paragraph.text.strip():
            text_parts.append(paragraph.text)

    has_tables = bool(doc.tables)
    for table in doc.tables:
        for row in table.rows:
            cells = [
                (cell.text or "").replace("\n", " ").strip()
                for cell in row.cells
            ]
            # Tab-separated preserves column structure for evidence
            # extractors that look for header→value patterns.
            if any(cells):
                text_parts.append("\t".join(cells))

    full_text = "\n".join(text_parts)

    return {
        "page_count": 1,
        "pages": [
            {
                "page_number": 1,
                "text": full_text,
                "has_tables": has_tables,
            }
        ],
        "is_scanned": False,
    }
