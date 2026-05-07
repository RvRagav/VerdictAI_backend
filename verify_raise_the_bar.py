"""Verification harness for the "raise the bar" upgrades.

Runs the four feature checks described in the task brief:

  1. Multi-format ingest — PNG, DOCX, and PDF upload + process.
  2. Explanation service — PASS / FAIL / REVIEW explanations populated.
  3. Reproducibility — end-to-end re-run matches byte-for-byte.
  4. PDF report — >5 pages, QR code present.

Each check is a standalone function that prints a one-line PASS/FAIL and
returns a boolean. The script exits 0 if all pass, 1 otherwise.
"""

from __future__ import annotations

import os
import sqlite3
import sys
import tempfile
import uuid
from datetime import datetime, timezone
from pathlib import Path


# Make `backend.*` importable when running this file directly.
_ROOT = Path(__file__).resolve().parent.parent
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))


# Keep LLM calls off — we rely on the rules branch for determinism.
os.environ.setdefault("LLM_DISABLED", "1")


GREEN = "\033[92m"
RED = "\033[91m"
YELLOW = "\033[93m"
RESET = "\033[0m"


def _pass(msg: str) -> bool:
    print(f"{GREEN}  PASS{RESET} {msg}")
    return True


def _fail(msg: str) -> bool:
    print(f"{RED}  FAIL{RESET} {msg}")
    return False


def _header(title: str) -> None:
    print()
    print(f"{YELLOW}{title}{RESET}")
    print("-" * len(title))


# ─── Fixtures ────────────────────────────────────────────────────────────


def _fresh_db() -> tuple[sqlite3.Connection, str, str]:
    """Create a temp DB with a seed tender and return (conn, tender_id, path)."""
    from backend.database.connection import get_db
    from backend.database.schema import create_tables

    path = tempfile.mktemp(suffix=".db")
    conn = get_db(path)
    create_tables(conn)

    tender_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    conn.execute(
        "INSERT INTO tenders (id, title, department, category, status, "
        "created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
        (tender_id, "Verification Tender",
         "CENTRAL RESERVE POLICE FORCE", "Works",
         "DOCUMENTS_UPLOADED", now, now),
    )
    conn.commit()
    return conn, tender_id, path


def _make_test_png(out_dir: Path) -> Path:
    """Create a simple PNG with legible text so OCR finds words."""
    from PIL import Image, ImageDraw, ImageFont

    path = out_dir / "test_certificate.png"
    # A4-ish dimensions at 150 DPI.
    img = Image.new("RGB", (1240, 1754), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype(
            "/System/Library/Fonts/Supplemental/Arial.ttf", 36
        )
    except OSError:
        font = ImageFont.load_default()

    lines = [
        "CHARTERED ACCOUNTANT CERTIFICATE",
        "",
        "This is to certify that M/s ApexGuard Technologies Pvt Ltd",
        "had an annual turnover of Rs. 18.45 Crore for FY 2023-24.",
        "",
        "GST: 07AAACA1234F1Z5",
        "PAN: AAACA1234F",
        "Valid up to 31-12-2026.",
    ]
    y = 120
    for line in lines:
        draw.text((80, y), line, fill="black", font=font)
        y += 60

    img.save(path)
    return path


def _make_test_docx(out_dir: Path) -> Path:
    """Create a DOCX with a paragraph and a table for parsing tests."""
    from docx import Document

    path = out_dir / "test_submission.docx"
    doc = Document()
    doc.add_heading("Bidder Submission — ApexGuard Technologies Pvt Ltd", 0)
    doc.add_paragraph(
        "We hereby submit our bid for the above tender. Our annual "
        "turnover for FY 2023-24 is Rs. 18.45 Crore."
    )
    doc.add_paragraph("GST: 07AAACA1234F1Z5   PAN: AAACA1234F")

    table = doc.add_table(rows=2, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Description"
    hdr[1].text = "Value"
    row = table.rows[1].cells
    row[0].text = "Turnover"
    row[1].text = "Rs. 18.45 Crore"

    doc.save(path)
    return path


# ─── Check 1: multi-format ingest ────────────────────────────────────────


def check_multi_format() -> bool:
    _header("1. Multi-format document ingest")

    from backend.layers.l1_document import process_document

    ok = True
    with tempfile.TemporaryDirectory() as tmpdir:
        out_dir = Path(tmpdir)

        # ── PNG ──
        conn, tender_id, db_path = _fresh_db()
        try:
            png = _make_test_png(out_dir)
            result = process_document(conn, tender_id, str(png), "certificate")
            conn.commit()

            if result["processing_status"] == "complete" \
                    and result["page_count"] == 1:
                _pass(
                    f"PNG processed: {result['page_count']} page, "
                    f"OCR conf {result['avg_ocr_confidence']:.2f}"
                )
            else:
                ok = _fail(f"PNG result: {result}") and ok

            # OCR must have produced some words
            word_count = conn.execute(
                "SELECT COUNT(*) c FROM word_objects "
                "WHERE page_id IN (SELECT id FROM pages "
                "WHERE document_id = ?)",
                (result["id"],),
            ).fetchone()["c"]
            if word_count > 5:
                _pass(f"PNG OCR produced {word_count} word objects")
            else:
                ok = _fail(f"PNG OCR produced only {word_count} words") and ok

            events = [
                e["event_type"] for e in conn.execute(
                    "SELECT event_type FROM audit_events "
                    "WHERE tender_id = ?",
                    (tender_id,),
                ).fetchall()
            ]
            if "document_received" in events and "ocr_completed" in events:
                _pass("PNG path logged document_received + ocr_completed")
            else:
                ok = _fail(f"PNG missing audit events: {events}") and ok
        finally:
            conn.close()
            if os.path.exists(db_path):
                os.unlink(db_path)

        # ── DOCX ──
        conn, tender_id, db_path = _fresh_db()
        try:
            docx = _make_test_docx(out_dir)
            result = process_document(conn, tender_id, str(docx),
                                      "bidder_submission")
            conn.commit()

            if result["processing_status"] == "complete":
                _pass("DOCX processed: digital text path")
            else:
                ok = _fail(f"DOCX result: {result}") and ok

            page = conn.execute(
                "SELECT raw_text FROM pages WHERE document_id = ?",
                (result["id"],),
            ).fetchone()
            if page and "18.45" in (page["raw_text"] or ""):
                _pass(
                    "DOCX text extracted (found turnover figure in raw_text)"
                )
            else:
                ok = _fail(
                    f"DOCX raw_text missing expected content: "
                    f"{page['raw_text'][:80] if page else None}"
                ) and ok
        finally:
            conn.close()
            if os.path.exists(db_path):
                os.unlink(db_path)

        # ── PDF (existing sample) ──
        conn, tender_id, db_path = _fresh_db()
        try:
            pdf = "backend/demo_data/sample_nit_crpf.pdf"
            if not os.path.exists(pdf):
                pdf = "backend/demo_data/sample_nit.pdf"
            if os.path.exists(pdf):
                result = process_document(conn, tender_id, pdf, "nit")
                conn.commit()
                if result["processing_status"] == "complete" \
                        and result["page_count"] >= 1:
                    _pass(
                        f"PDF still processes: {result['page_count']} pages"
                    )
                else:
                    ok = _fail(f"PDF result: {result}") and ok
            else:
                _fail("Demo PDF missing — skipping PDF regression check")
        finally:
            conn.close()
            if os.path.exists(db_path):
                os.unlink(db_path)

    return ok


# ─── Check 2: explanation service ────────────────────────────────────────


def check_explanation_service() -> bool:
    _header("2. Explanation service")

    from backend.services.explanation_service import build_explanation

    ok = True

    # PASS — numeric threshold.
    pass_criterion = {
        "criterion_type": "numeric_threshold",
        "criterion_text": "Average annual turnover >= Rs. 15 Crore",
        "source_clause_ref": "4.1(a)",
        "threshold_value": {"value": 15, "unit": "Cr", "rupees": 150000000},
        "is_mandatory": True,
        "gfr_override_permitted": False,
    }
    pass_evidence = {
        "value": {
            "amount": 184500000, "raw_value": 18.45,
            "unit": "crore", "fiscal_year": "2023-24",
        },
        "source_page_number": 1,
        "source_filename": "ca_certificate.pdf",
        "ocr_confidence": 0.92,
        "extraction_confidence": 0.91,
        "evaluation_method": "deterministic",
        "entity_match_flag": False,
    }
    exp = build_explanation("PASS", pass_criterion, pass_evidence,
                            "auto_commit")
    problems = []
    if not exp["headline"]:
        problems.append("empty headline")
    if "18.45" not in exp["headline"] and "Rs." not in exp["headline"]:
        problems.append("headline missing amount")
    if not exp["detail"]:
        problems.append("empty detail")
    if not exp["facts"]:
        problems.append("empty facts")
    if not exp["source_reference"]:
        problems.append("empty source_reference")
    if problems:
        ok = _fail(f"PASS numeric: {problems}\n    {exp}") and ok
    else:
        _pass(f"PASS numeric: '{exp['headline']}'")

    # FAIL — entity mismatch.
    fail_criterion = {
        "criterion_type": "categorical_presence",
        "criterion_text": "Bidder identity matches submitted documents",
        "source_clause_ref": "4.2",
        "threshold_value": {"required": True},
        "is_mandatory": True,
        "gfr_override_permitted": False,
    }
    fail_evidence = {
        "value": {
            "found": True,
            "registration_number": "07AAACA1234F1Z5",
            "certificate_type": "gst",
            "is_valid": True,
            "validity_date": "2026-12-31",
        },
        "source_page_number": 3,
        "source_filename": "gst_cert.pdf",
        "ocr_confidence": 0.88,
        "extraction_confidence": 0.60,
        "evaluation_method": "deterministic",
        "entity_match_flag": True,
        "entity_match_result": {
            "registered_name": "ApexGuard Technologies Pvt Ltd",
            "extracted_name": "ApexGuard Group International Ltd",
            "mismatch_type": "parent_company",
            "llm_classification": "parent_company",
            "llm_reasoning":
                "Extracted name is the parent-group entity; fraud vector.",
        },
    }
    exp2 = build_explanation("FAIL", fail_criterion, fail_evidence,
                             "mandatory_review")
    problems = []
    if not exp2["headline"]:
        problems.append("empty headline")
    if not any("Entity mismatch" in f or "entity mismatch" in f.lower()
               for f in exp2["facts"]):
        problems.append("facts missing entity mismatch note")
    if "Officer confirmation required" not in exp2["next_action"] \
            and "officer" not in exp2["next_action"].lower():
        problems.append("next_action doesn't prompt officer")
    if problems:
        ok = _fail(f"FAIL entity: {problems}\n    {exp2}") and ok
    else:
        _pass(f"FAIL entity: '{exp2['headline']}'")

    # REVIEW — stamp obscuration.
    rev_criterion = {
        "criterion_type": "numeric_threshold",
        "criterion_text": "Turnover >= Rs. 15 Crore",
        "source_clause_ref": "4.1(a)",
        "threshold_value": {"value": 15, "unit": "Cr", "rupees": 150000000},
        "is_mandatory": True,
        "gfr_override_permitted": False,
    }
    rev_evidence = {
        "value": {
            "amount": 72000000, "raw_value": 7.2,
            "unit": "crore", "fiscal_year": "2023-24",
        },
        "source_page_number": 1,
        "source_filename": "ca_certificate.pdf",
        "ocr_confidence": 0.41,
        "extraction_confidence": 0.41,
        "evaluation_method": "deterministic",
        "entity_match_flag": False,
        "stamp_regions": [
            {"x": 100, "y": 200, "w": 300, "h": 80, "area": 24000},
            {"x": 500, "y": 800, "w": 250, "h": 60, "area": 15000},
        ],
    }
    exp3 = build_explanation("REVIEW", rev_criterion, rev_evidence,
                             "hitl_review")
    problems = []
    if not any("stamp" in f.lower() for f in exp3["facts"]):
        problems.append("facts missing stamp region note")
    if not exp3["confidence_note"]:
        problems.append("empty confidence_note")
    if problems:
        ok = _fail(f"REVIEW stamp: {problems}\n    {exp3}") and ok
    else:
        _pass(f"REVIEW stamp: '{exp3['headline']}'")

    return ok


# ─── Check 3: reproducibility re-run ─────────────────────────────────────


def _setup_full_tender(conn: sqlite3.Connection) -> str:
    """Create a tender with bidders, criteria, and approved statuses.

    Used by the reproducibility and PDF report checks — the seed data
    doesn't attach bidders to tenders, so we build a minimal but
    complete tender here.
    """
    import json
    tender_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    conn.execute(
        "INSERT INTO tenders (id, title, department, category, status, "
        "created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
        (tender_id, "Supply of Security Equipment — Verification Run",
         "CENTRAL RESERVE POLICE FORCE", "security_equipment",
         "EVALUATING", now, now),
    )

    # Two bidders so the summary table has variety.
    bidders = [
        ("ApexGuard Technologies Pvt Ltd", "AAACA1234F"),
        ("SecureVision Systems Pvt Ltd", "AABCS5678G"),
    ]
    bidder_ids: list[str] = []
    for name, pan in bidders:
        bid_id = str(uuid.uuid4())
        bidder_ids.append(bid_id)
        conn.execute(
            "INSERT INTO bidders (id, tender_id, company_name, pan_number, "
            "registration_number, status, debarment_status) "
            "VALUES (?, ?, ?, ?, ?, ?, ?)",
            (bid_id, tender_id, name, pan, "REG-" + pan[:6],
             "pending", "clear"),
        )

        # One bidder submission per bidder so evidence extraction
        # has documents to look at.
        doc_id = str(uuid.uuid4())
        conn.execute(
            "INSERT INTO documents (id, tender_id, bidder_id, doc_type, "
            "filename, file_path, sha256_hash, page_count, "
            "avg_ocr_confidence, upload_timestamp, processing_status) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (doc_id, tender_id, bid_id, "bidder_submission",
             f"{name.replace(' ', '_')}_submission.pdf",
             f"/tmp/{bid_id}.pdf",
             "0" * 64, 1, 0.9, now, "complete"),
        )
        page_id = str(uuid.uuid4())
        conn.execute(
            "INSERT INTO pages (id, document_id, page_number, raw_text, "
            "ocr_confidence) VALUES (?, ?, ?, ?, ?)",
            (
                page_id, doc_id, 1,
                f"M/s {name}\n"
                f"PAN: {pan}   GST: 07{pan}1Z5\n"
                "Annual turnover for FY 2023-24: Rs. 18.45 Crore.\n"
                "Valid up to 31-12-2026.\n"
                "Completed 4 similar supply orders in the last 5 years.",
                0.9,
            ),
        )

    # Criteria covering all three deterministic types (qualitative
    # excluded so reproducibility stays byte-identical without LLM).
    criteria_data = [
        (
            "Average annual turnover >= Rs. 10 Crore in last 3 FYs",
            "numeric_threshold",
            json.dumps({"value": 10, "unit": "Cr",
                        "rupees": 100000000, "period": "3 years",
                        "label": "turnover"}),
            1, 0, "Clause 4.1(a)",
        ),
        (
            "Valid GST registration certificate",
            "categorical_presence",
            json.dumps({"document": "GST", "required": True}),
            1, 0, "Clause 4.1(b)",
        ),
        (
            "Minimum 2 similar supply orders in last 5 years",
            "temporal_recency",
            json.dumps({"count": 2, "period": "5 years"}),
            0, 1, "Clause 4.2(a)",
        ),
    ]
    for (text, ctype, threshold, mandatory, override, clause) in criteria_data:
        conn.execute(
            "INSERT INTO criteria (id, tender_id, criterion_text, "
            "criterion_type, threshold_value, gfr_override_permitted, "
            "is_mandatory, source_clause_ref, status) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (str(uuid.uuid4()), tender_id, text, ctype, threshold,
             override, mandatory, clause, "approved"),
        )

    conn.commit()
    return tender_id


def check_reproducibility() -> bool:
    _header("3. Reproducibility")

    from backend.database.connection import get_db
    from backend.database.schema import create_tables
    from backend.database.seed import seed_demo_data
    from backend.layers.l3_evidence import extract_evidence  # noqa: F401
    from backend.layers.l4_evaluation import evaluate_all_bidders
    from backend.services.reproducibility import reproduce_evaluation

    ok = True
    db_path = tempfile.mktemp(suffix=".db")
    conn = get_db(db_path)
    create_tables(conn)
    seed_demo_data(conn)

    try:
        tender_id = _setup_full_tender(conn)

        results = evaluate_all_bidders(conn, tender_id)
        conn.commit()

        if not results:
            return _fail("evaluate_all_bidders returned no results")

        eval_count = conn.execute(
            "SELECT COUNT(*) c FROM evaluations WHERE tender_id = ?",
            (tender_id,),
        ).fetchone()["c"]
        _pass(
            f"Evaluated {len(results)} bidder(s), "
            f"{eval_count} total evaluations"
        )

        repro = reproduce_evaluation(conn, tender_id)
        conn.commit()

        if repro["match"]:
            _pass(
                f"Re-run matches ({repro['matches']}/"
                f"{repro['total_compared']} evaluations), "
                f"byte_identical={repro['byte_identical_excluding_timestamps']}"
            )
        else:
            ok = _fail(
                f"Re-run DIFFS: {len(repro['diffs'])} diffs — "
                f"first: {repro['diffs'][:1]}"
            )

        if repro["reproducibility_hash_original"] == \
                repro["reproducibility_hash_reproduced"]:
            _pass(
                f"Reproducibility hashes match: "
                f"{repro['reproducibility_hash_original'][:16]}..."
            )
        else:
            ok = _fail(
                "Reproducibility hashes differ: "
                f"orig={repro['reproducibility_hash_original'][:16]}... "
                f"repro={repro['reproducibility_hash_reproduced'][:16]}..."
            ) and ok
    finally:
        conn.close()
        if os.path.exists(db_path):
            os.unlink(db_path)

    return ok


# ─── Check 4: PDF report ────────────────────────────────────────────────


def check_pdf_report() -> bool:
    _header("4. Print-ready PDF report")

    from backend.database.connection import get_db
    from backend.database.schema import create_tables
    from backend.database.seed import seed_demo_data
    from backend.layers.l4_evaluation import evaluate_all_bidders
    from backend.services.report_service import generate_report

    ok = True
    db_path = tempfile.mktemp(suffix=".db")
    conn = get_db(db_path)
    create_tables(conn)
    seed_demo_data(conn)

    try:
        tender_id = _setup_full_tender(conn)
        evaluate_all_bidders(conn, tender_id)
        conn.commit()

        report = generate_report(
            conn=conn,
            tender_id=tender_id,
            officer_id="officer-verdictai-test",
        )
        conn.commit()

        pdf_path = Path(report["download_path"])
        if pdf_path.suffix != ".pdf" or not pdf_path.exists():
            return _fail(
                f"PDF was not generated: download_path={pdf_path}"
            )
        size_kb = pdf_path.stat().st_size / 1024
        _pass(f"PDF generated at {pdf_path} ({size_kb:.1f} KB)")

        # Count pages from the PDF itself. PyPDF2 is already a dependency.
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(pdf_path))
            page_count = len(reader.pages)
        except Exception as exc:
            return _fail(f"Could not read PDF pages: {exc}")

        if page_count > 5:
            _pass(f"PDF has {page_count} pages (> 5 required)")
        else:
            ok = _fail(f"PDF has only {page_count} pages (need > 5)") and ok

        # Extract text from the signature page and look for the hash.
        try:
            last_page_text = reader.pages[-1].extract_text() or ""
        except Exception:
            last_page_text = ""

        hash_hex = report["sha256_hash"]
        # Hash is rendered with spaces every 16 chars; look for a prefix.
        if hash_hex[:16] in last_page_text.replace(" ", ""):
            _pass(
                f"Signature page contains audit hash "
                f"({hash_hex[:16]}...)"
            )
        else:
            ok = _fail(
                "Audit hash not found on signature page "
                f"(hash={hash_hex[:32]}...)"
            ) and ok

        # QR code detection: the PDF should contain at least one /Image
        # XObject (the QR bitmap). We check resource streams directly.
        qr_found = False
        for page in reader.pages:
            resources = page.get("/Resources")
            if not resources:
                continue
            xobj = resources.get("/XObject") if hasattr(
                resources, "get") else None
            if xobj:
                try:
                    for _name, ref in xobj.items():
                        obj = ref.get_object()
                        if obj.get("/Subtype") == "/Image":
                            qr_found = True
                            break
                except Exception:
                    continue
            if qr_found:
                break
        if qr_found:
            _pass("PDF contains at least one embedded image (QR code)")
        else:
            ok = _fail("No embedded image detected in PDF (QR missing?)") and ok

    finally:
        conn.close()
        if os.path.exists(db_path):
            os.unlink(db_path)

    return ok


# ─── Check 5: audit chain API shape ──────────────────────────────────────


def check_audit_chain_api() -> bool:
    """Backend endpoints the new Audit Trail + Reproduce UIs consume."""
    _header("5. Audit chain API")

    from backend.database.connection import get_db
    from backend.database.schema import create_tables
    from backend.database.seed import seed_demo_data
    from backend.layers.l4_evaluation import evaluate_all_bidders
    from backend.layers.l5_audit import get_audit_trail

    ok = True
    db_path = tempfile.mktemp(suffix=".db")
    conn = get_db(db_path)
    create_tables(conn)
    seed_demo_data(conn)

    try:
        tender_id = _setup_full_tender(conn)
        evaluate_all_bidders(conn, tender_id)
        conn.commit()

        # The API wraps get_audit_trail + adds prev_hash/entry_hash fields.
        # Exercise get_audit_trail directly and assert the hash fields are
        # populated so the frontend has something to render.
        trail = get_audit_trail(conn, tender_id)
        if not trail:
            return _fail("No audit events found after evaluation")
        _pass(f"Audit trail has {len(trail)} events")

        with_hashes = [
            e for e in trail if e.get("entry_hash") and e.get("prev_hash") is not None
        ]
        if len(with_hashes) == len(trail):
            _pass("Every event carries prev_hash + entry_hash")
        else:
            ok = _fail(
                f"{len(trail) - len(with_hashes)}/{len(trail)} "
                "events missing hash fields"
            ) and ok

        # First entry's prev_hash must be the genesis (all zeros).
        first = trail[0]
        if first["prev_hash"] == "0" * 64:
            _pass("Genesis event uses zero prev_hash")
        else:
            ok = _fail(f"Unexpected prev_hash on first event: {first['prev_hash']}") and ok
    finally:
        conn.close()
        if os.path.exists(db_path):
            os.unlink(db_path)

    return ok


# ─── Entry point ─────────────────────────────────────────────────────────


def main() -> int:
    # Silence pdfplumber / PIL INFO warnings that fire when the
    # verification tender uses synthetic paths that don't exist on disk.
    import logging
    logging.getLogger("pdfminer").setLevel(logging.CRITICAL)
    logging.getLogger("backend.services.evidence_extractor").setLevel(
        logging.CRITICAL
    )

    print(f"{YELLOW}VerdictAI — Raise-the-Bar Verification{RESET}")
    results = {
        "Multi-format ingest": check_multi_format(),
        "Explanation service": check_explanation_service(),
        "Reproducibility": check_reproducibility(),
        "PDF report": check_pdf_report(),
        "Audit chain API": check_audit_chain_api(),
    }

    print()
    print(f"{YELLOW}Summary{RESET}")
    print("-" * 7)
    all_ok = True
    for name, ok in results.items():
        status = f"{GREEN}PASS{RESET}" if ok else f"{RED}FAIL{RESET}"
        print(f"  {status}  {name}")
        all_ok = all_ok and ok

    return 0 if all_ok else 1


if __name__ == "__main__":
    sys.exit(main())
